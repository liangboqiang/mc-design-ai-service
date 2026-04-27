from __future__ import annotations

import json
import re
import time
import uuid
from pathlib import Path
from typing import Any


REPORT_STATE = "design_reports.json"


class DesignReportToolbox:
    """Embedded design report toolbox migrated from the old report MCP toolbox."""

    toolbox_name = "design_report"
    tags = ("external", "report", "document")

    def __init__(self, workspace_root: Path | None = None):
        self.workspace_root = workspace_root.resolve() if workspace_root else None
        self.runtime = None

    def bind_runtime(self, runtime, tool_lookup=None) -> None:  # noqa: ANN001
        self.runtime = runtime

    def spawn(self, workspace_root: Path) -> "DesignReportToolbox":
        return DesignReportToolbox(workspace_root=workspace_root)

    def executors(self):
        return {
            "design_report.create_report": self._exec_create_report,
            "design_report.get_report_detail": self._exec_get_report_detail,
            "design_report.update_report": self._exec_update_report,
            "design_report.list_reports": self._exec_list_reports,
            "design_report.export_report": self._exec_export_report,
            "design_report.create_image": self._exec_create_image,
        }

    def _exec_create_report(self, args: dict):
        template_path = str(args["template_path"])
        conversation_id = str(args.get("conversation_id") or self.runtime.settings.conversation_id)
        user_id = str(args.get("user_id") or self.runtime.settings.user_id)
        state = self._load_all()
        report_id = f"rpt_{uuid.uuid4().hex[:10]}"
        slots = self._extract_slots(template_path)
        report = {
            "report_id": report_id,
            "user_id": user_id,
            "conversation_id": conversation_id,
            "template_path": template_path,
            "report_status": "draft",
            "slots": slots,
            "outline": self._outline(slots),
            "created_at": time.time(),
            "updated_at": time.time(),
        }
        state[report_id] = report
        self._save_all(state)
        allowed = {"slot_id", "type", "name", "section", "prompt", "fixed", "fill"}
        return json.dumps(
            {
                "ok": True,
                "data": {
                    "report_id": report_id,
                    "report_status": "draft",
                    "slots": [{k: v for k, v in slot.items() if k in allowed} for slot in slots],
                },
                "message": "报告创建成功",
            },
            ensure_ascii=False,
            indent=2,
        )

    def _exec_get_report_detail(self, args: dict):
        report = self._get_report(str(args["report_id"]))
        return json.dumps({"ok": True, "data": self._detail(report), "message": "获取报告详情成功"}, ensure_ascii=False, indent=2)

    def _exec_update_report(self, args: dict):
        report_id = str(args["report_id"])
        content = args.get("content") or []
        state = self._load_all()
        report = state.get(report_id)
        if not report:
            return json.dumps({"ok": False, "data": None, "message": f"报告不存在: {report_id}"}, ensure_ascii=False, indent=2)
        slot_map = {slot["slot_id"]: slot for slot in report.get("slots", [])}
        for item in content:
            if not isinstance(item, dict):
                continue
            slot_id = item.get("slot_id")
            if slot_id not in slot_map:
                continue
            slot = slot_map[slot_id]
            for key in ("text", "paragraph", "table", "image"):
                if key in item:
                    slot[key] = item[key]
                    slot["fill"] = True
                    break
        report["updated_at"] = time.time()
        report["validation"] = self._validate(report)
        report["report_status"] = "ready" if not report["validation"]["missing_slots"] else "draft"
        state[report_id] = report
        self._save_all(state)
        return json.dumps(
            {
                "ok": True,
                "data": {
                    "report_id": report_id,
                    "report_status": report["report_status"],
                    "outline": report.get("outline", []),
                    "validation": report.get("validation", {}),
                },
                "message": "报告更新成功",
            },
            ensure_ascii=False,
            indent=2,
        )

    def _exec_list_reports(self, args: dict):
        user_id = str(args.get("user_id") or self.runtime.settings.user_id)
        conversation_id = str(args.get("conversation_id") or "")
        rows = []
        for report in self._load_all().values():
            if report.get("user_id") != user_id:
                continue
            if conversation_id and report.get("conversation_id") != conversation_id:
                continue
            slots = report.get("slots", [])
            rows.append(
                {
                    "report_id": report.get("report_id"),
                    "report_status": report.get("report_status"),
                    "template_path": report.get("template_path"),
                    "conversation_id": report.get("conversation_id"),
                    "slot_total": len(slots),
                    "slot_filled": len([s for s in slots if s.get("fill")]),
                    "updated_at": report.get("updated_at"),
                }
            )
        rows.sort(key=lambda x: float(x.get("updated_at") or 0), reverse=True)
        return json.dumps({"ok": True, "data": {"reports": rows}, "message": "获取报告列表成功"}, ensure_ascii=False, indent=2)

    def _exec_export_report(self, args: dict):
        report = self._get_report(str(args["report_id"]))
        export_path = str(args.get("export_path") or "").strip()
        if not export_path:
            template = Path(str(report.get("template_path") or "report_template.docx"))
            export_path = str(template.with_name(template.stem + "_exported.docx"))
        target = self._safe_output_path(export_path)
        target.parent.mkdir(parents=True, exist_ok=True)

        if str(report.get("template_path", "")).lower().endswith(".docx"):
            self._export_docx(report, target)
        else:
            target.write_text(self._render_plain(report), encoding="utf-8")

        report["export_path"] = str(target)
        report["updated_at"] = time.time()
        state = self._load_all()
        state[report["report_id"]] = report
        self._save_all(state)
        return json.dumps(
            {"ok": True, "data": {"export_path": str(target)}, "message": "报告导出成功"},
            ensure_ascii=False,
            indent=2,
        )

    def _exec_create_image(self, args: dict):
        image_path = str(args.get("image_path") or "")
        nx_tool = self._lookup_nx_create_image()
        if nx_tool:
            return nx_tool({"user_id": args.get("user_id") or self.runtime.settings.user_id, "filePath": image_path})
        return json.dumps(
            {
                "ok": False,
                "data": None,
                "message": "NX CreateImage executor is not installed. Enable nx toolbox or call nx.CreateImage directly.",
            },
            ensure_ascii=False,
            indent=2,
        )

    def _load_all(self) -> dict[str, Any]:
        return self.runtime.session.read_state_json(REPORT_STATE, {}) or {}

    def _save_all(self, state: dict[str, Any]) -> None:
        self.runtime.session.write_state_json(REPORT_STATE, state)

    def _get_report(self, report_id: str) -> dict[str, Any]:
        report = self._load_all().get(report_id)
        if not report:
            raise ValueError(f"报告不存在: {report_id}")
        return report

    def _extract_slots(self, template_path: str) -> list[dict[str, Any]]:
        path = Path(template_path)
        text = ""
        try:
            if path.exists() and path.suffix.lower() in {".txt", ".md", ".json"}:
                text = path.read_text(encoding="utf-8")
            elif path.exists() and path.suffix.lower() == ".docx":
                from docx import Document

                doc = Document(str(path))
                text = "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            text = ""

        anchors = sorted(set(re.findall(r"\{\{\s*([A-Za-z0-9_\\-]+)\s*\}\}", text)))
        if not anchors:
            anchors = ["summary", "parameters", "conclusion"]
        slots = []
        for index, anchor in enumerate(anchors, start=1):
            slots.append(
                {
                    "slot_id": f"SLOT_{index:03d}",
                    "type": "text",
                    "name": anchor,
                    "section": "default",
                    "prompt": f"请填写 {anchor}",
                    "anchor": "{{" + anchor + "}}",
                    "fixed": False,
                    "fill": False,
                }
            )
        return slots

    @staticmethod
    def _outline(slots: list[dict[str, Any]]) -> list[dict[str, Any]]:
        return [{"section": slot.get("section", "default"), "slot_id": slot["slot_id"], "name": slot.get("name")} for slot in slots]

    @staticmethod
    def _validate(report: dict[str, Any]) -> dict[str, Any]:
        missing = [slot["slot_id"] for slot in report.get("slots", []) if not slot.get("fill")]
        return {"missing_slots": missing, "ready": not missing}

    def _detail(self, report: dict[str, Any]) -> dict[str, Any]:
        report = dict(report)
        report["validation"] = self._validate(report)
        return report

    def _render_plain(self, report: dict[str, Any]) -> str:
        lines = [f"# Design Report {report.get('report_id')}", ""]
        for slot in report.get("slots", []):
            lines.append(f"## {slot.get('name') or slot.get('slot_id')}")
            lines.append(self._slot_value_to_text(slot))
            lines.append("")
        return "\n".join(lines)

    def _export_docx(self, report: dict[str, Any], target: Path) -> None:
        try:
            from docx import Document
        except Exception:
            target.with_suffix(".md").write_text(self._render_plain(report), encoding="utf-8")
            return

        template_path = Path(str(report.get("template_path") or ""))
        if template_path.exists():
            doc = Document(str(template_path))
        else:
            doc = Document()
            doc.add_heading("Design Report", level=1)
            for slot in report.get("slots", []):
                doc.add_paragraph(f"{slot.get('name')}: {slot.get('anchor')}")
        replacements = {slot.get("anchor"): self._slot_value_to_text(slot) for slot in report.get("slots", []) if slot.get("anchor")}
        for paragraph in doc.paragraphs:
            text = paragraph.text or ""
            changed = False
            for anchor, value in replacements.items():
                if anchor and anchor in text:
                    text = text.replace(anchor, value)
                    changed = True
            if changed:
                for run in paragraph.runs:
                    run.text = ""
                paragraph.add_run(text)
        doc.save(str(target))

    @staticmethod
    def _slot_value_to_text(slot: dict[str, Any]) -> str:
        for key in ("text", "paragraph", "table", "image"):
            value = slot.get(key)
            if not value:
                continue
            if isinstance(value, dict):
                if key == "text":
                    return str(value.get("content") or "")
                if key == "paragraph":
                    if isinstance(value.get("blocks"), list):
                        return "\n".join(str(x) for x in value["blocks"])
                    return str(value.get("content") or "")
                if key == "table":
                    rows = value.get("rows") or []
                    return "\n".join("\t".join(str(x) for x in row) if isinstance(row, list) else str(row) for row in rows)
                if key == "image":
                    return str(value.get("path") or value.get("nx_path") or "")
            return str(value)
        return ""

    def _safe_output_path(self, raw: str) -> Path:
        path = Path(raw)
        if path.is_absolute():
            return path
        if self.workspace_root is None:
            raise ValueError("workspace not bound")
        return (self.workspace_root / path).resolve()

    def _lookup_nx_create_image(self):
        if self.runtime is None:
            return None
        return self.runtime.runtime_state.tool_registry.get("nx.CreateImage").executor if "nx.CreateImage" in self.runtime.runtime_state.tool_registry else None
