from __future__ import annotations

import csv
import difflib
import io
import json
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

from .catalogs import DESIGN_RULES, FIELD_REQUIREMENTS, KNOWLEDGE_CATALOG, PART_ALIASES, TOOL_CATALOG


class TextOpsToolbox:
    toolbox_name = "textops"
    tags = ("text", "local_compute", "data_processing")

    def __init__(self, workspace_root: Path | None = None):
        self.workspace_root = workspace_root.resolve() if workspace_root else None
        self.runtime = None

    def bind_runtime(self, runtime, tool_lookup=None) -> None:  # noqa: ANN001
        self.runtime = runtime

    def spawn(self, workspace_root: Path) -> "TextOpsToolbox":
        return TextOpsToolbox(workspace_root=workspace_root)

    def executors(self):
        return {
            "textops.search": self._exec_textops_search,
            "textops.preview_replace": self._exec_textops_preview_replace,
            "textops.parse_text_from_url": self._exec_parse_text_from_url,
            "textops.parse_docx_from_url": self._exec_parse_docx_from_url,
            "textops.design_rule_lookup": self._exec_design_rule_lookup,
            "textops.parameter_check": self._exec_parameter_check,
            "textops.estimate_action_cost": self._exec_estimate_action_cost,
            "textops.tool_catalog_query": self._exec_tool_catalog_query,
            "textops.knowledge_catalog_query": self._exec_knowledge_catalog_query,
        }

    def _exec_textops_search(self, args: dict):
        return self._search(args["path"], args["query"])

    def _exec_textops_preview_replace(self, args: dict):
        return self._preview_replace(args["path"], args["old_text"], args["new_text"])

    def _exec_parse_text_from_url(self, args: dict):
        return json.dumps(self._parse_text(args["url"], args.get("encoding", "utf-8")), ensure_ascii=False, indent=2)

    def _exec_parse_docx_from_url(self, args: dict):
        return json.dumps(self._parse_docx(args["url"]), ensure_ascii=False, indent=2)

    def _exec_design_rule_lookup(self, args: dict):
        part_type = self._part_type(args.get("part_type", ""))
        keywords = [str(x).lower() for x in args.get("keywords", []) or []]
        rule_ref = str(args.get("rule_ref") or "").strip()
        rows = []
        for rule in DESIGN_RULES:
            if part_type and rule["part_type"] not in {"*", part_type}:
                continue
            if rule_ref and rule["rule_ref"] != rule_ref:
                continue
            hay = json.dumps(rule, ensure_ascii=False).lower()
            if keywords and not all(k in hay for k in keywords):
                continue
            rows.append(rule)
        return json.dumps({"ok": True, "data": {"rules": rows}, "message": "设计规则检索完成"}, ensure_ascii=False, indent=2)

    def _exec_parameter_check(self, args: dict):
        part_type = self._part_type(args["part_type"])
        facts = args.get("facts") or []
        fact_map = {}
        for item in facts:
            if isinstance(item, dict):
                if "field" in item and "value" in item:
                    fact_map[str(item["field"])] = item.get("value")
                else:
                    fact_map.update(item)
        issues = []
        for req in sorted([r for r in FIELD_REQUIREMENTS if r["part_type"] == part_type], key=lambda x: x.get("order", 0)):
            value = fact_map.get(req["field"])
            if req.get("required") and (value is None or value == ""):
                issues.append({"type": "missing_required", "field": req["field"], "severity": "high", "ask_text": req.get("ask_text"), "rule_ref": req.get("rule_ref")})
        for rule in DESIGN_RULES:
            if rule["part_type"] not in {"*", part_type}:
                continue
            rr = rule["rule_ref"]
            if rr == "ROD-101" and not self._positive(fact_map.get("center_distance")):
                issues.append({"type": "rule_violation", "rule_ref": rr, "severity": rule["severity"], "message": rule["rule_text"]})
            if rr == "ROD-102" and not self._positive(fact_map.get("small_end_diameter")):
                issues.append({"type": "rule_violation", "rule_ref": rr, "severity": rule["severity"], "message": rule["rule_text"]})
            if rr == "ROD-103" and self._num(fact_map.get("big_end_diameter")) is not None and self._num(fact_map.get("small_end_diameter")) is not None and self._num(fact_map.get("big_end_diameter")) <= self._num(fact_map.get("small_end_diameter")):
                issues.append({"type": "rule_violation", "rule_ref": rr, "severity": rule["severity"], "message": rule["rule_text"]})
            if rr == "CRK-101" and not self._positive(fact_map.get("stroke")):
                issues.append({"type": "rule_violation", "rule_ref": rr, "severity": rule["severity"], "message": rule["rule_text"]})
            if rr == "CAM-101" and not self._positive_int(fact_map.get("cam_count")):
                issues.append({"type": "rule_violation", "rule_ref": rr, "severity": rule["severity"], "message": rule["rule_text"]})
            if rr == "CAM-103" and self._num(fact_map.get("duration_angle")) is not None and not (0 <= self._num(fact_map.get("duration_angle")) <= 360):
                issues.append({"type": "rule_violation", "rule_ref": rr, "severity": rule["severity"], "message": rule["rule_text"]})
        return json.dumps({"ok": True, "data": {"issues": issues}, "message": "参数校核完成"}, ensure_ascii=False, indent=2)

    def _exec_estimate_action_cost(self, args: dict):
        actions = args.get("actions") or []
        policy = args.get("policy") or {}
        weights = {"ask_user": 1.0, "call_tool": 2.0, "write_file": 1.5, "spawn_subagent": 4.0, "default": 1.0}
        weights.update({k: float(v) for k, v in policy.items() if isinstance(v, (int, float))})
        items = []
        total = 0.0
        for action in actions:
            kind = str(action.get("type") or action.get("action") or "default") if isinstance(action, dict) else "default"
            cost = float(weights.get(kind, weights["default"]))
            total += cost
            items.append({"action": action, "cost": cost})
        return json.dumps({"ok": True, "data": {"cost_items": items, "total_cost": total}, "message": "动作成本估算完成"}, ensure_ascii=False, indent=2)

    def _exec_tool_catalog_query(self, args: dict):
        role = str(args.get("tool_role") or "")
        keywords = [str(x).lower() for x in args.get("keywords", []) or []]
        rows = []
        for item in TOOL_CATALOG:
            if role and item.get("tool_role") != role:
                continue
            hay = json.dumps(item, ensure_ascii=False).lower()
            if keywords and not all(k in hay for k in keywords):
                continue
            rows.append(item)
        return json.dumps({"ok": True, "data": {"tools": rows}, "message": "工具目录查询完成"}, ensure_ascii=False, indent=2)

    def _exec_knowledge_catalog_query(self, args: dict):
        ktype = str(args.get("knowledge_type") or "")
        keywords = [str(x).lower() for x in args.get("keywords", []) or []]
        rows = []
        for item in KNOWLEDGE_CATALOG:
            if ktype and item.get("knowledge_type") != ktype:
                continue
            hay = json.dumps(item, ensure_ascii=False).lower()
            if keywords and not all(k in hay for k in keywords):
                continue
            rows.append(item)
        return json.dumps({"ok": True, "data": {"items": rows}, "message": "知识目录查询完成"}, ensure_ascii=False, indent=2)

    def _safe_path(self, raw: str) -> Path:
        if self.workspace_root is None:
            raise ValueError("TextOpsToolbox workspace not bound yet.")
        path = (self.workspace_root / raw).resolve()
        if not path.is_relative_to(self.workspace_root):
            raise ValueError(f"Path escapes workspace: {raw}")
        return path

    def _read_url_or_path(self, url: str, encoding: str = "utf-8") -> tuple[str, dict[str, Any]]:
        parsed = urlparse(url)
        if parsed.scheme in {"http", "https"}:
            import requests

            response = requests.get(url, timeout=30)
            response.raise_for_status()
            return response.content.decode(encoding, errors="replace"), {"source": url, "bytes": len(response.content)}
        path = Path(url.replace("file://", ""))
        if not path.is_absolute():
            path = self._safe_path(str(path))
        raw = path.read_bytes()
        return raw.decode(encoding, errors="replace"), {"source": str(path), "bytes": len(raw)}

    def _parse_text(self, url: str, encoding: str = "utf-8") -> dict[str, Any]:
        text, meta = self._read_url_or_path(url, encoding)
        suffix = Path(urlparse(url).path).suffix.lower()
        rows = []
        if suffix == ".json":
            try:
                rows = json.loads(text)
            except Exception:
                rows = []
        elif suffix == ".csv":
            rows = list(csv.reader(io.StringIO(text)))
        return {"ok": True, "data": {"text": text, "preview": text[:2000], "rows": rows, "meta": meta}, "message": "文本解析完成"}

    def _parse_docx(self, url: str) -> dict[str, Any]:
        parsed = urlparse(url)
        if parsed.scheme in {"http", "https"}:
            import requests

            response = requests.get(url, timeout=30)
            response.raise_for_status()
            payload = io.BytesIO(response.content)
        else:
            path = Path(url.replace("file://", ""))
            if not path.is_absolute():
                path = self._safe_path(str(path))
            payload = str(path)
        try:
            from docx import Document
        except Exception as exc:
            return {"ok": False, "data": None, "message": f"python-docx 不可用: {exc}"}
        doc = Document(payload)
        text = "\n".join(p.text for p in doc.paragraphs)
        tables = []
        for table in doc.tables:
            tables.append([[cell.text for cell in row.cells] for row in table.rows])
        return {"ok": True, "data": {"text": text, "preview": text[:2000], "tables": tables, "meta": {"source": url}}, "message": "docx解析完成"}

    def _search(self, path: str, query: str) -> str:
        lines = self._safe_path(path).read_text(encoding="utf-8").splitlines()
        matches = [f"{index}: {line}" for index, line in enumerate(lines, start=1) if query in line]
        return "\n".join(matches) or "(no matches)"

    def _preview_replace(self, path: str, old_text: str, new_text: str) -> str:
        target = self._safe_path(path)
        before = target.read_text(encoding="utf-8")
        after = before.replace(old_text, new_text, 1)
        return "\n".join(difflib.unified_diff(before.splitlines(), after.splitlines(), fromfile=path, tofile=f"{path} (preview)", lineterm=""))

    @staticmethod
    def _part_type(raw: str) -> str:
        raw = str(raw or "").strip()
        return PART_ALIASES.get(raw, raw)

    @staticmethod
    def _num(value: Any) -> float | None:
        try:
            return float(value)
        except Exception:
            return None

    @classmethod
    def _positive(cls, value: Any) -> bool:
        n = cls._num(value)
        return n is not None and n > 0

    @classmethod
    def _positive_int(cls, value: Any) -> bool:
        n = cls._num(value)
        return n is not None and n > 0 and int(n) == n
